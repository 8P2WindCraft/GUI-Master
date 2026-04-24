import os
import re
import copy
import shutil
from docx import Document
from docx.shared import Cm
import pandas as pd
from openpyxl import load_workbook

def svg_to_png_bytes(svg_path, log_callback=None, dateipfad=None, scale=3):
    """
    Konvertiert eine SVG-Datei in PNG-Bytes (für python-docx).
    Nutzt PySide6 (Qt), da es bereits für die GUI verwendet wird.
    """
    log_prefix = f"[{os.path.basename(dateipfad)}]" if dateipfad else "[SVG]"
    def _log(msg):
        if log_callback:
            log_callback(f"{log_prefix} {msg}")

    _log(f"Starte SVG-Konvertierung mit PySide6 für: {svg_path}")

    if not os.path.exists(svg_path):
        _log(f"FEHLER: SVG-Datei nicht gefunden am Pfad: {svg_path}")
        return None

    try:
        from PySide6.QtSvg import QSvgRenderer
        from PySide6.QtGui import QImage, QPainter
        from PySide6.QtCore import QByteArray, QBuffer, QIODevice, QSize
        from io import BytesIO

        with open(svg_path, 'rb') as f:
            svg_data = f.read()

        renderer = QSvgRenderer(QByteArray(svg_data))
        if not renderer.isValid():
            _log(f"FEHLER: PySide6 konnte die SVG-Datei nicht laden: {svg_path}. Ist es eine gültige SVG?")
            return None

        size = renderer.defaultSize()
        if size.isEmpty():
            _log(f"WARNUNG: SVG hat keine Standardgröße. Verwende 300x150 als Basis.")
            size.setWidth(300)
            size.setHeight(150)
        
        # Skalierungsfaktor für höhere Auflösung (z.B. 3 für 3x so viele Pixel in jeder Dimension)
        scaled_size = QSize(size.width() * scale, size.height() * scale)

        image = QImage(scaled_size, QImage.Format_ARGB32)
        image.fill(0x00000000)  # Transparenter Hintergrund

        painter = QPainter(image)
        renderer.render(painter)
        painter.end()

        byte_array = QByteArray()
        buffer = QBuffer(byte_array)
        buffer.open(QIODevice.OpenModeFlag.WriteOnly)
        image.save(buffer, "PNG")
        buffer.close()

        if byte_array.isEmpty():
            _log("FEHLER: Konvertierung resultierte in einem leeren PNG-Buffer.")
            return None

        _log(f"SVG erfolgreich nach PNG konvertiert. Buffer-Größe: {len(byte_array)} bytes.")
        return byte_array.data()

    except ImportError:
        _log("FEHLER: 'PySide6' nicht gefunden. Dies sollte nicht passieren, da es für die GUI benötigt wird.")
        return None
    except Exception as e:
        import traceback
        _log(f"FATALER FEHLER bei SVG-Konvertierung mit PySide6: {e}")
        _log(f"Traceback: {traceback.format_exc()}")
        return None

def is_file_locked(filepath):
    """Überprüft, ob eine Datei bereits von einem anderen Prozess geöffnet ist.
    
    Args:
        filepath: Pfad zur zu überprüfenden Datei
        
    Returns:
        bool: True, wenn die Datei gesperrt ist, sonst False
    """
    import errno
    
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Datei nicht gefunden: {filepath}")
    
    try:
        # Versuche, die Datei exklusiv zu öffnen
        fd = os.open(filepath, os.O_WRONLY | os.O_CREAT | os.O_EXCL)
        # Wenn erfolgreich, sofort wieder schließen
        os.close(fd)
        # Die temporäre Datei wieder löschen
        os.remove(filepath)
        return False
    except OSError as e:
        if e.errno == errno.EEXIST:  # Datei existiert bereits
            return False
        if e.errno in (errno.EACCES, errno.EAGAIN, errno.ETXTBSY):  # Datei ist gesperrt
            return True
        # Andere Fehler weiterwerfen
        raise

def lade_excel_daten(pfad, header_row=3, log_callback=None):
    """Lädt Daten aus einer Excel-Datei und gibt sie als Liste von Dictionaries zurück.
    
    Args:
        pfad: Pfad zur Excel-Datei
        header_row: Zeilennummer der Überschriften (1-basiert)
        log_callback: Optionaler Callback für Log-Nachrichten
        
    Returns:
        tuple: (Liste der Datensätze als Dictionaries, Dictionary mit Standardwerten)
        
    Raises:
        FileNotFoundError: Wenn die Datei nicht gefunden wird
        PermissionError: Wenn die Datei bereits geöffnet ist
        Exception: Bei anderen Fehlern
    """
    def _log(msg):
        if log_callback:
            log_callback(msg)
        print(msg)
    
    _log(f"Lade Excel-Daten aus: {pfad}")
    
    # Überprüfen, ob die Datei existiert
    if not os.path.isfile(pfad):
        error_msg = f"FEHLER: Excel-Datei nicht gefunden: {pfad}"
        _log(error_msg)
        raise FileNotFoundError(error_msg)
    
    # Überprüfen, ob die Datei bereits geöffnet ist
    if is_file_locked(pfad):
        error_msg = f"FEHLER: Die Datei '{os.path.basename(pfad)}' ist bereits in einem anderen Programm geöffnet. Bitte schließen Sie die Datei und versuchen Sie es erneut."
        _log(error_msg)
        raise PermissionError(error_msg)
    
    try:
        import pandas as pd
        
        # Versuche, die Excel-Datei zu lesen
        _log(f"Öffne Excel-Datei: {pfad}")
        xls = pd.ExcelFile(pfad)
        
        # Erste Tabelle auswählen
        sheet = xls.sheet_names[0]
        _log(f"Lese Arbeitsblatt: {sheet}")
        
        # Daten einlesen
        df = pd.read_excel(pfad, sheet_name=sheet, header=header_row-1)
        
        # Leere Zeilen entfernen
        df = df.dropna(how='all')
        
        # Spaltennamen bereinigen
        df.columns = [str(col).strip() for col in df.columns]
        
        _log(f"Excel-Header: {list(df.columns)}")
        
        # In Liste von Dictionaries umwandeln
        datensaetze = []
        for _, row in df.iterrows():
            ds = row.to_dict()
            seriennummer = ds.get('anlage_seriennummer')
            if seriennummer is not None and str(seriennummer).strip() != '':
                datensaetze.append(ds)
        
        _log(f"Erfolgreich {len(datensaetze)} Datensätze geladen.")
        
        # Fallback-Markierungen aus zweitem Arbeitsblatt laden (falls vorhanden)
        fallback_marken = {}
        try:
            from openpyxl import load_workbook
            
            _log("Versuche, Textmarken aus zweitem Arbeitsblatt zu laden...")
            wb = load_workbook(pfad, data_only=True, read_only=True)
            available_sheets = wb.sheetnames
            
            if len(available_sheets) > 1:
                textmarken_ws = wb[available_sheets[1]]
                for row in textmarken_ws.iter_rows(min_row=2, values_only=True):
                    if row is None or len(row) < 3:
                        continue
                    _, marke, wert, *rest = row + (None, None, None)  # Sicherstellen, dass genug Werte vorhanden sind
                    if marke is not None and str(marke).strip() != '':
                        fallback_marken[str(marke).strip()] = str(wert) if wert is not None else ""
                
                _log(f"Erfolgreich {len(fallback_marken)} Fallback-Textmarken geladen.")
            else:
                _log("Kein zweites Arbeitsblatt für Textmarken gefunden.")
                
        except Exception as e:
            _log(f"Warnung: Konnte Textmarken nicht laden: {str(e)}")
        
        return datensaetze, fallback_marken
        
    except PermissionError as e:
        # Falls die Datei zwischenzeitlich gesperrt wurde
        if "Permission denied" in str(e) or "[Errno 13]" in str(e):
            error_msg = f"FEHLER: Keine Berechtigung zum Lesen der Datei '{os.path.basename(pfad)}'. Stellen Sie sicher, dass die Datei nicht in einem anderen Programm geöffnet ist."
            _log(error_msg)
            raise PermissionError(error_msg) from e
        raise
    except Exception as e:
        error_msg = f"FEHLER beim Lesen der Excel-Datei {pfad}: {str(e)}"
        _log(error_msg)
        raise Exception(error_msg) from e
    return datensaetze, fallback_marken

# Hilfsfunktion zum sicheren Setzen von Text
def _set_run_text(run, text):
    """Setzt den Text eines Runs und kümmert sich um Zeilenumbrüche."""
    if run is None or text is None:
        return
    
    # Text in saubere Zeichenkette umwandeln, um Probleme mit non-string Werten zu vermeiden
    text = str(text)

    run.text = '' # Bestehenden Text löschen
    
    # Zeilenumbrüche manuell als <w:br/> einfügen
    parts = text.split('\n')
    for i, part in enumerate(parts):
        run.add_text(part)
        if i < len(parts) - 1:
            run.add_break()

def ersetze_textmarken(doc, datensatz, fallback, dateipfad=None, log_callback=None, bilder_ordner=None, excel_df=None, svg_scale=3):
    """
    Ersetzt alle {{...}}-Platzhalter robust in Paragraphen und Tabellen.
    - Phase 1: Ersetzt Bild- und QR-Code-Platzhalter, die in einem eigenen Paragraphen stehen.
    - Phase 2: Ersetzt Text-Platzhalter im gesamten Dokument (Paragraphen und Tabellen).
    """
    import os
    from docx.text.run import Run
    from datetime import datetime
    from docx.shared import Cm
    import qrcode
    from io import BytesIO

    fehlende = set()
    pattern = re.compile(r"{{(.*?)}}")

    # Phase 1: Bilder und QR-Codes ersetzen (nur wenn der Paragraph ausschließlich den Platzhalter enthält)
    for para in doc.paragraphs[:]:
        full_text = para.text.strip()
        # Verwenden von fullmatch, um sicherzustellen, dass der gesamte Paragraph nur den Platzhalter enthält
        match = pattern.fullmatch(full_text)
        if not match:
            continue

        key = match.group(1).strip()
        is_image = key.endswith('_img')
        is_qr = key.endswith('_link')

        if not (is_image or is_qr):
            continue
        
        # Hole Wert aus den Hauptdaten (Blatt 1)
        wert = datensatz.get(key)
        
        # Verwende Fallback-Wert aus Blatt 2, wenn Wert in Blatt 1 None oder leer ist
        if wert is None or (isinstance(wert, str) and not wert.strip()):
            wert = fallback.get(key)
        
        # Wenn immer noch kein Wert, überspringe
        if wert is None or not str(wert).strip():
            _log(f"WARNUNG: Kein Wert für Bild/QR '{key}' gefunden. Platzhalter wird entfernt.", dateipfad)
            # Entferne den leeren Paragraphen, um Lücken im Dokument zu vermeiden
            p = para._element
            p.getparent().remove(p)
            para._p = para._element = None
            continue
        
        wert = str(wert).strip()
        
        # Lösche den Platzhalter-Paragraphen
        # Wichtig: Der neue Inhalt wird an der gleichen Stelle eingefügt
        p_element = para._element
        p_parent = p_element.getparent()
        p_index = p_parent.index(p_element)
        p_parent.remove(p_element)
        
        # Erstelle einen neuen Paragraphen für das Bild/QR
        new_para = p_parent.insert_paragraph_before("", p_parent[p_index] if p_index < len(p_parent) else None)
        run = new_para.add_run()
        
        try:
            if is_image:
                # Bildpfad zusammenbauen
                bild_pfad = os.path.join(bilder_ordner, wert) if bilder_ordner else wert
                _log(f"INFO: Ersetze Bild-Platzhalter '{key}' mit Bild: {bild_pfad}", dateipfad)
                if not os.path.exists(bild_pfad):
                    fehlende.add(key)
                    _log(f"FEHLER: Bilddatei nicht gefunden: {bild_pfad}", dateipfad)
                    run.text = f"[FEHLER: Bild nicht gefunden: {wert}]"
                    run.font.color.rgb = (255, 0, 0)
                else:
                    # Bild einfügen
                    if bild_pfad.lower().endswith('.svg'):
                        png_bytes = svg_to_png_bytes(bild_pfad, log_callback, dateipfad, scale=svg_scale)
                        if png_bytes:
                            run.add_picture(BytesIO(png_bytes), width=Cm(15))
                        else:
                            fehlende.add(key)
                            _log(f"FEHLER: SVG-Konvertierung fehlgeschlagen für: {bild_pfad}", dateipfad)
                            run.text = f"[FEHLER: SVG-Konvertierung fehlgeschlagen: {wert}]"
                            run.font.color.rgb = (255, 0, 0)
                    else:
                        run.add_picture(bild_pfad, width=Cm(15))
            
            elif is_qr:
                # QR-Code generieren
                _log(f"INFO: Ersetze QR-Platzhalter '{key}' mit QR-Code für: {wert}", dateipfad)
                qr_img = qrcode.make(wert)
                img_bytes = BytesIO()
                qr_img.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                run.add_picture(img_bytes, width=Cm(4))
        except Exception as e:
            _log(f"FEHLER beim Verarbeiten von Bild/QR '{key}': {e}", dateipfad)
            run.text = f"[FEHLER: {e}]"
            run.font.color.rgb = (255, 0, 0)
            
    # Phase 2: Text-Platzhalter im gesamten Dokument ersetzen (Paragraphen und Tabellen)
    all_runs = []
    for para in doc.paragraphs:
        all_runs.extend(para.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_runs.extend(para.runs)

    # Puffern der zusammenhängenden Platzhalter-Runs
    i = 0
    while i < len(all_runs):
        run = all_runs[i]
        # Prüfen, ob der Run-Text einen Platzhalter-Anfang enthält
        if '{{' in run.text:
            # Sammle alle zusammenhängenden Runs, die einen vollständigen Platzhalter bilden könnten
            # Beispiel: `{{` in run1, `platzhalter` in run2, `}}` in run3
            collected_runs = [run]
            full_text = run.text
            j = i + 1
            # Suche nach dem schließenden Teil des Platzhalters
            while '}}' not in full_text and j < len(all_runs):
                next_run = all_runs[j]
                full_text += next_run.text
                collected_runs.append(next_run)
                j += 1

            # Führe Ersetzungen für den gesammelten Text durch
            if '}}' in full_text:
                # Finde alle Platzhalter im gesammelten Text
                for match in pattern.finditer(full_text):
                    key = match.group(1).strip()
                    
                    # Wert aus Hauptdaten oder Fallback holen
                    wert = datensatz.get(key)
                    if wert is None:
                        wert = fallback.get(key)
                    
                    # Wenn kein Wert gefunden wurde, markieren
                    if wert is None:
                        fehlende.add(key)
                        wert = f"[{key}_NICHT_GEFUNDEN]"
                    
                    # Ersetze den Platzhalter im Text
                    full_text = full_text.replace(match.group(0), str(wert))
                
                # Den neuen, ersetzten Text in den ersten Run schreiben
                # und die restlichen (jetzt überflüssigen) Runs leeren.
                first_run = collected_runs[0]
                _set_run_text(first_run, full_text)

                for k in range(1, len(collected_runs)):
                    collected_runs[k].text = ''

            # Springe zum nächsten Run nach der verarbeiteten Gruppe
            i = j
        else:
            i += 1
            
    if fehlende:
        _log(f"WARNUNG: Folgende Textmarken wurden nicht in der Excel-Datei gefunden und konnten nicht ersetzt werden: {sorted(list(fehlende))}", dateipfad)

    return doc, fehlende

def finde_table_config(doc, prefix):
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == f"{prefix}_TABLE_CONFIG":
            if idx + 1 < len(doc.paragraphs):
                return doc.paragraphs[idx + 1].text.strip()
    return None

def ersetze_table(doc, excel_path, config, table_marker):
    if not config:
        return

    spalten_part, *optionen = config.split(";")
    spalten_idx = [int(s.strip()) - 1 for s in spalten_part.split(",")]
    options = {k.strip(): int(v.strip()) for k, v in (o.split("=") for o in optionen)}

    df = pd.read_excel(excel_path, sheet_name="Daten", header=options["ÜZ"] - 1)
    df = df.iloc[options["START"] - 1:]
    df = df.iloc[:, spalten_idx]

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == table_marker:
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for j, col in enumerate(df.columns):
                hdr_cells[j].text = str(col)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for j, val in enumerate(row):
                    row_cells[j].text = str(val)
            p.clear()
            break

def verarbeite_vorlagen(vorlagen_ordner, export_ordner, excel_path, progress_callback=None, status_callback=None, header_row=3, bilder_ordner=None, log_callback=None, svg_scale=3, categories=None):
    # Standard-Kategorien, falls keine übergeben wurden
    if categories is None:
        categories = {
            'b_': 'Beschilderung',
            'ba_': 'Betriebsanweisung',
            'p_': 'Pläne',
            'a_': 'Allgemein',
            'hv_': 'Havarieplan'
        }
    def _log(msg):
        print(msg)
        if log_callback:
            log_callback(msg)
    _log(f"verarbeite_vorlagen aufgerufen mit vorlagen_ordner={vorlagen_ordner}, export_ordner={export_ordner}, excel_path={excel_path}, header_row={header_row}")
    from datetime import datetime
    # Zeitstempel generieren
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    base_export = os.path.join(export_ordner, ts)
    # Basis-Exportordner erstellen
    os.makedirs(base_export, exist_ok=True)

    print("Lade Excel-Daten...")
    datensaetze, fallback = lade_excel_daten(excel_path, header_row=header_row)
    print(f"Excel-Datensätze: {len(datensaetze)}, Fallbacks: {len(fallback)}")

    # Fortschritt initialisieren
    total_steps = 0
    allgemein_path = os.path.join(vorlagen_ordner, "Allgemein")
    anlage_path = os.path.join(vorlagen_ordner, "Anlagen")
    if os.path.isdir(allgemein_path):
        allgemein_files = [f for f in os.listdir(allgemein_path) if f.endswith(".docx")]
        total_steps += len(allgemein_files)
    else:
        allgemein_files = []
    if os.path.isdir(anlage_path):
        anlagen_files = [f for f in os.listdir(anlage_path) if f.endswith(".docx")]
        total_steps += len(anlagen_files) * len(datensaetze)
    else:
        anlagen_files = []
    if total_steps == 0:
        total_steps = 1
    current_step = 0
    if progress_callback:
        print(f"Fortschritt: {current_step}/{total_steps}")
        progress_callback(current_step, total_steps)

    # Hilfsfunktion zum Bestimmen der Kategorie
    def get_category_and_clean_name(filename):
        """Bestimmt die Kategorie und bereinigt den Dateinamen
        
        Args:
            filename: Der ursprüngliche Dateiname
            
        Returns:
            tuple: (Zielordner, bereinigter Dateiname)
        """
        # Dateiname ohne Pfad
        base_name = os.path.basename(filename)
        
        # Prüfe auf bekannte Präfixe
        for prefix, category in categories.items():
            if base_name.lower().startswith(prefix.lower()):
                # Entferne das Präfix aus dem Dateinamen
                clean_name = base_name[len(prefix):]
                return category, clean_name
        
        # Standard: Keine Kategorie gefunden, in Allgemein/Anlagen belassen
        return None, base_name
        return None, filename

    # Verarbeite alle Dokumente aus dem Vorlagenordner
    print(f"Prüfe Vorlagen: {vorlagen_ordner}")
    if os.path.isdir(vorlagen_ordner):
        # Alle .docx Dateien im Hauptverzeichnis verarbeiten
        all_files = [f for f in os.listdir(vorlagen_ordner) if f.lower().endswith('.docx')]
        print(f"Gefundene Dateien: {all_files}")
        
        for idx, file in enumerate(all_files, 1):
            if status_callback:
                status_callback(f"Verarbeite: {file} ({idx}/{len(all_files)})")
            
            # Kategorie bestimmen und Dateinamen bereinigen
            category, clean_filename = get_category_and_clean_name(file)
            
            # Zielverzeichnis erstellen
            if category:
                target_dir = os.path.join(base_export, category, str(idx))
            else:
                target_dir = os.path.join(base_export, str(idx))
            
            os.makedirs(target_dir, exist_ok=True)
            out_path = os.path.join(target_dir, clean_filename)
            _log(f"Datei {file} wird in Kategorie {category} exportiert als {clean_filename}")
            
            # Dokument kopieren (oder verarbeiten, falls benötigt)
            try:
                shutil.copy2(os.path.join(vorlagen_ordner, file), out_path)
                _log(f"Erfolgreich kopiert nach: {out_path}")
            except Exception as e:
                _log(f"FEHLER beim Kopieren von {file}: {str(e)}")
            
            current_step += 1
            if progress_callback:
                progress_callback(current_step, total_steps)

    # Anlagendokumente verarbeiten (falls vorhanden)
    if datensaetze and os.path.isdir(vorlagen_ordner):
        # Nach Unterordnern für Anlagendokumente suchen
        for item in os.listdir(vorlagen_ordner):
            item_path = os.path.join(vorlagen_ordner, item)
            if os.path.isdir(item_path):
                _log(f"Verarbeite Unterordner: {item}")
                
                # Alle .docx Dateien im Unterordner verarbeiten
                anlagen_files = [f for f in os.listdir(item_path) if f.lower().endswith('.docx')]
                
                for filename in anlagen_files:
                    try:
                        doc_path = os.path.join(item_path, filename)
                        doc = Document(doc_path)
                
                        # Bestimme Kategorie und bereinige den Dateinamen
                        category, clean_name = get_category_and_clean_name(filename)
                        
                        # Wenn eine Kategorie gefunden wurde, erstelle den entsprechenden Ordner
                        if category:
                            category_path = os.path.join(base_export, category)
                            os.makedirs(category_path, exist_ok=True)
                            _log(f"Kategorie erkannt: {filename} -> {category}")
                        
                        for datensatz in datensaetze:
                            try:
                                # Erstelle eine Kopie des Dokuments für jeden Datensatz
                                doc_copy = Document()
                                for element in doc.element.body:
                                    doc_copy.element.body.append(copy.deepcopy(element))
                                
                                # Ersetze die Platzhalter
                                ersetze_textmarken(
                                    doc=doc_copy,
                                    datensatz=datensatz,
                                    fallback=fallback,
                                    dateipfad=os.path.join(item_path, filename),
                                    log_callback=log_callback,
                                    bilder_ordner=bilder_ordner,
                                    excel_df=pd.DataFrame(datensaetze),
                                    svg_scale=svg_scale
                                )
                                
                                # Speichere die Datei mit dem Anlagennamen
                                anlagenname = datensatz.get('Anlagenname', 'Unbekannt').replace('/', '_')
                                base_name = os.path.splitext(clean_name)[0]  # Ohne Präfix
                                export_name = f"{base_name}_{anlagenname}.docx"
                                
                                # Bestimme den Zielpfad basierend auf der Kategorie
                                if category:
                                    export_path = os.path.join(category_path, export_name)
                                else:
                                    export_path = os.path.join(anlagen_export, export_name)
                                    
                                doc_copy.save(export_path)
                                _log(f"Erstellt: {export_path}")
                                
                            except Exception as e:
                                _log(f"Fehler bei der Verarbeitung von {filename} für {datensatz.get('Anlagenname', 'Unbekannt')}: {e}")
                            
                            current_step += 1
                            if progress_callback:
                                progress_callback(current_step, total_steps)
                    
                    except Exception as e:
                        _log(f"Fehler beim Öffnen von {filename}: {e}")
                        current_step += len(datensaetze)  # Überspringe alle Schritte für diese Datei
                        if progress_callback:
                            progress_callback(current_step, total_steps)

    else:
        print(f"Anlagen-Ordner nicht gefunden: {anlage_path}")
        if status_callback:
            status_callback(f"WARNUNG: Anlagen-Ordner nicht gefunden: {anlage_path}")
    # Fortschritt auf 100% setzen
    if progress_callback:
        progress_callback(total_steps, total_steps)

# GUI für Dokument- und Verzeichnisauswahl

import json
import os

def run_with_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk
    import threading

    SETTINGS_FILE = os.path.join(os.path.dirname(__file__), "settings.json")

    def load_settings():
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def save_settings():
        data = {
            "excel": excel_var.get(),
            "vorlagen": vorlagen_var.get(),
            "export": export_var.get()
        }
        try:
            with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f)
        except Exception:
            pass

    def update_progress(current, total):
        progress_bar['maximum'] = total
        progress_bar['value'] = current
        progress_label_var.set(f"Fortschritt: {current} / {total}")
        root.update_idletasks()

    def update_status(text):
        status_var.set(text)
        root.update_idletasks()

    def set_and_save(var, value):
        var.set(value)
        save_settings()

    def start_processing():
        excel = excel_var.get()
        eingang = vorlagen_var.get()
        export = export_var.get()
        status_var.set("")
        progress_bar['value'] = 0
        progress_label_var.set("")
        save_settings()
        if not excel or not eingang or not export:
            status_var.set("Bitte alle Felder ausfüllen!")
            return
        def worker():
            try:
                verarbeite_vorlagen(
                    eingang, export, excel,
                    progress_callback=lambda c, t: root.after(0, update_progress, c, t),
                    status_callback=lambda text: root.after(0, update_status, text)
                )
                root.after(0, update_status, "Fertig!")
            except Exception as e:
                root.after(0, update_status, f"Fehler: {e}")
        threading.Thread(target=worker, daemon=True).start()

    root = tk.Tk()
    root.title("Bericht Generator")

    settings = load_settings()

    # Grid-Konfiguration für dynamische Anpassung
    root.columnconfigure(0, weight=0)
    root.columnconfigure(1, weight=1)
    root.columnconfigure(2, weight=0)
    for i in range(7):
        root.rowconfigure(i, weight=0)

    tk.Label(root, text="Excel-Datei:").grid(row=0, column=0, sticky="e", padx=2, pady=2)
    excel_var = tk.StringVar(value=settings.get("excel", ""))
    tk.Entry(root, textvariable=excel_var).grid(row=0, column=1, sticky="ew", padx=2, pady=2)
    tk.Button(root, text="Durchsuchen...", command=lambda: set_and_save(excel_var, filedialog.askopenfilename(title="Excel-Datei auswählen", filetypes=[("Excel Dateien", "*.xlsx *.xls")]))).grid(row=0, column=2, sticky="ew", padx=2, pady=2)

    tk.Label(root, text="Vorlagen-Verzeichnis:").grid(row=1, column=0, sticky="e", padx=2, pady=2)
    vorlagen_var = tk.StringVar(value=settings.get("vorlagen", ""))
    tk.Entry(root, textvariable=vorlagen_var).grid(row=1, column=1, sticky="ew", padx=2, pady=2)
    tk.Button(root, text="Durchsuchen...", command=lambda: set_and_save(vorlagen_var, filedialog.askdirectory(title="Vorlagen-Verzeichnis auswählen"))).grid(row=1, column=2, sticky="ew", padx=2, pady=2)

    tk.Label(root, text="Export-Verzeichnis:").grid(row=2, column=0, sticky="e", padx=2, pady=2)
    export_var = tk.StringVar(value=settings.get("export", ""))
    tk.Entry(root, textvariable=export_var).grid(row=2, column=1, sticky="ew", padx=2, pady=2)
    tk.Button(root, text="Durchsuchen...", command=lambda: set_and_save(export_var, filedialog.askdirectory(title="Export-Verzeichnis auswählen"))).grid(row=2, column=2, sticky="ew", padx=2, pady=2)

    status_var = tk.StringVar()
    tk.Label(root, textvariable=status_var, fg="red").grid(row=3, column=0, columnspan=3, sticky="ew", padx=2, pady=2)

    progress_label_var = tk.StringVar()
    tk.Label(root, textvariable=progress_label_var).grid(row=4, column=0, columnspan=3, sticky="ew", padx=2, pady=2)

    progress_bar = ttk.Progressbar(root, orient="horizontal", mode="determinate")
    progress_bar.grid(row=5, column=0, columnspan=3, sticky="ew", padx=2, pady=5)

    tk.Button(root, text="Start", command=start_processing, bg="#4CAF50", fg="white").grid(row=6, column=0, columnspan=3, pady=10, sticky="ew")

    root.minsize(500, 250)
    root.mainloop()

if __name__ == "__main__":
    run_with_gui()

